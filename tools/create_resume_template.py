from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "后端开发简历模板.docx"

BLUE = "0F4B8F"
DEEP_BLUE = "073763"
WHITE = "FFFFFF"
TEXT = "1F2933"
MUTED = "5B6673"
LIGHT_LINE = "D7DEE8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        data = kwargs.get(edge)
        if data is None:
            continue
        tag = "start" if edge == "left" else "end" if edge == "right" else edge
        element = tc_borders.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tc_borders.append(element)
        for key in ["sz", "val", "color", "space"]:
            if key in data:
                element.set(qn(f"w:{key}"), str(data[key]))


def set_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def run_text(paragraph, text, size=9, bold=False, color=TEXT, font="Microsoft YaHei"):
    run = paragraph.add_run(text)
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def paragraph(cell, text="", size=9, bold=False, color=TEXT, align=None, before=0, after=2, line=1.05):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align:
        p.alignment = align
    if text:
        run_text(p, text, size=size, bold=bold, color=color)
    return p


def clear_cell(cell):
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)


def section_title(cell, title, icon=""):
    p = paragraph(cell, "", after=4)
    if icon:
        run_text(p, icon + "  ", size=11, bold=True, color=BLUE)
    r = run_text(p, title, size=12, bold=True, color=WHITE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), BLUE)
    p_pr.append(shd)
    return r


def side_title(cell, title):
    p = paragraph(cell, title, size=12, bold=True, color=WHITE, before=8, after=5)
    p.paragraph_format.keep_with_next = True


def bullet(cell, text, size=8.2, color=WHITE, level=0):
    p = paragraph(cell, after=2)
    p.paragraph_format.left_indent = Cm(0.25 + level * 0.2)
    p.paragraph_format.first_line_indent = Cm(-0.12)
    run_text(p, "• ", size=size, color=color, bold=True)
    run_text(p, text, size=size, color=color)


def right_bullet(cell, text, bold_prefix=None):
    p = paragraph(cell, after=2)
    p.paragraph_format.left_indent = Cm(0.34)
    p.paragraph_format.first_line_indent = Cm(-0.18)
    run_text(p, "• ", size=8.5, color=TEXT, bold=True)
    if bold_prefix and text.startswith(bold_prefix):
        run_text(p, bold_prefix, size=8.5, bold=True)
        run_text(p, text[len(bold_prefix):], size=8.5)
    else:
        run_text(p, text, size=8.5)


def job_row(cell, date, org, role):
    t = cell.add_table(rows=1, cols=3)
    t.autofit = False
    widths = [3.8, 7.0, 3.5]
    for i, w in enumerate(widths):
        set_width(t.cell(0, i), w)
        set_cell_margins(t.cell(0, i), 0, 0, 0, 0)
        set_cell_border(t.cell(0, i), top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
        clear_cell(t.cell(0, i))
    paragraph(t.cell(0, 0), date, size=9, bold=True)
    paragraph(t.cell(0, 1), org, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(t.cell(0, 2), role, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.9)
    section.bottom_margin = Cm(0.9)
    section.left_margin = Cm(1.15)
    section.right_margin = Cm(1.15)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(8.8)

    layout = doc.add_table(rows=1, cols=2)
    layout.autofit = False
    left = layout.cell(0, 0)
    right = layout.cell(0, 1)
    set_width(left, 6.1)
    set_width(right, 12.6)
    for c in (left, right):
        clear_cell(c)
        set_cell_border(c, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(left, DEEP_BLUE)
    set_cell_margins(left, top=260, start=260, bottom=180, end=260)
    set_cell_margins(right, top=80, start=300, bottom=80, end=90)

    photo = left.add_table(rows=1, cols=1)
    photo.autofit = False
    pc = photo.cell(0, 0)
    clear_cell(pc)
    set_width(pc, 3.1)
    set_cell_margins(pc, top=520, bottom=520, start=80, end=80)
    set_cell_shading(pc, WHITE)
    set_cell_border(pc, top={"val": "single", "sz": 8, "color": "FFFFFF"}, bottom={"val": "single", "sz": 8, "color": "FFFFFF"}, left={"val": "single", "sz": 8, "color": "FFFFFF"}, right={"val": "single", "sz": 8, "color": "FFFFFF"})
    paragraph(pc, "一寸照片", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    photo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph(left, "姓名", size=24, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=2)
    paragraph(left, "求职意向：后端开发", size=12, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    side_title(left, "基本信息")
    for item in ["出生年月：YYYY年MM月", "手机号码：13800000000", "微信：your_wechat", "邮箱：yourname@email.com", "政治面貌：共青团员", "最高学历：本科"]:
        paragraph(left, item, size=9.2, bold=True, color=WHITE, after=3)
    side_title(left, "获奖证书")
    for item in ["第十五届蓝桥杯国赛三等奖", "第十五届蓝桥杯省赛一等奖", "软件测试大赛省级优秀奖", "全国大学英语六级 CET-6", "华为云 / 阿里云相关认证"]:
        bullet(left, item)
    side_title(left, "荣誉称号 & 奖学金")
    for item in ["2023-2024 学年校级奖学金", "优秀学生干部 / 三好学生", "创新创业竞赛优秀奖", "优秀共青团员"]:
        bullet(left, item)

    paragraph(right, "个人简历", size=22, bold=True, color=BLUE, after=0)
    paragraph(right, "PERSONAL RESUME", size=16, bold=True, color=BLUE, after=6)
    section_title(right, "教育背景", "▣")
    job_row(right, "2022.09-2026.06", "XX大学（211）", "软件工程专业")
    right_bullet(right, "主修课程：数据结构、操作系统、计算机组成原理、计算机网络、Linux、算法分析与设计、数据库原理、J2EE开发框架")
    right_bullet(right, "GPA：3.99/5.00（前10%）")

    section_title(right, "实习经历", "▣")
    job_row(right, "2024.11-2025.01", "飞跃智能物联", "JAVA开发")
    paragraph(right, "岗位职责：负责项目的研发，包括前后端代码编写、嵌入式开发部门对接及联调。", size=8.5, color=MUTED, after=2)
    right_bullet(right, "业务代码编写：根据需求利用 WEB + ORM 后端框架编写接口与管理端功能。", "业务代码编写：")
    right_bullet(right, "配置相关服务：维护开发环境配置，使用 Nginx / Tomcat 完成部署联调。", "配置相关服务：")
    right_bullet(right, "实现硬件通信：利用 MQTT 消息队列完成设备间通信与状态同步。", "实现硬件通信：")
    right_bullet(right, "线程调度管理：使用线程池及异步任务提升任务处理效率。", "线程调度管理：")
    job_row(right, "2025.05-2025.08", "拼多多（暑期实习）", "服务端开发")
    paragraph(right, "岗位职责：负责商城业务迭代、接口开发、数据处理与线上查询链路优化。", size=8.5, color=MUTED, after=2)
    right_bullet(right, "业务需求开发：评审需求并完成 RPC 接口、数据模型和服务注册功能。", "业务需求开发：")
    right_bullet(right, "配置数据源：按业务逻辑拆分冷热数据，接入 MySQL / HBase / Hive 等。", "配置数据源：")
    right_bullet(right, "慢查询治理优化：分析日志与索引，降低接口响应时间。", "慢查询治理优化：")

    section_title(right, "项目经历", "▣")
    job_row(right, "2025.01-2025.03", "字节跳动青训营抖音商城", "后端开发")
    paragraph(right, "项目介绍：围绕电商交易链路完成商品、购物车、订单、支付等核心模块。", size=8.5, color=MUTED, after=2)
    paragraph(right, "技术栈：Go、Hertz、Kitex、Gorm、MySQL、Redis、Consul、Docker", size=8.5, bold=True, after=2)
    right_bullet(right, "使用 Go 语言与 Hertz 框架完成接口开发，保证模块边界清晰。")
    right_bullet(right, "使用 Kitex 构建 RPC 服务调用链路，降低系统耦合。")
    right_bullet(right, "使用 Docker 完成开发环境部署，便于联调与交付。")

    section_title(right, "技能评价", "▣")
    right_bullet(right, "熟悉 Java / Go 后端开发，了解 Spring Boot、微服务与常见中间件。")
    right_bullet(right, "熟悉 MySQL 索引、事务、慢查询分析及 Redis 缓存设计。")
    right_bullet(right, "具备良好的需求理解、问题定位、文档编写和跨团队沟通能力。")

    doc.save(OUT)


if __name__ == "__main__":
    main()
